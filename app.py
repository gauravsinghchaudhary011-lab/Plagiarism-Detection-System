from flask import Flask, render_template, request, flash
from backend.detection import detect_plagiarism
from backend.preprocess import preprocess_text
import fitz  # PyMuPDF
from docx import Document
import pytesseract
from PIL import Image
import io

app = Flask(__name__)
app.secret_key = 'pds-hacker-secret-key-2024-change-in-prod'


def parse_file(file) -> str:
    """Parse TXT, PDF, DOCX, image to text."""
    filename = (file.filename or "").lower()
    file_content = file.read()

    if filename.endswith('.txt'):
        return file_content.decode('utf-8', errors='ignore')

    if filename.endswith('.pdf'):
        doc = fitz.open(stream=file_content, filetype="pdf")
        texts = []
        try:
            for page in doc:
                # `text` is usually more reliable than the default
                page_text = page.get_text("text") or ""
                if page_text and page_text.strip():
                    texts.append(page_text.strip())
        finally:
            doc.close()
        return "\n".join(texts).strip()

    if filename.endswith('.docx'):
        doc = Document(io.BytesIO(file_content))

        parts = []

        # paragraphs
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                parts.append(para.text.strip())

        # tables (often contain important content)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        parts.append(cell.text.strip())

        return "\n".join(parts).strip()

    if any(img_ext in filename for img_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']):
        image = Image.open(io.BytesIO(file_content))
        text = pytesseract.image_to_string(image)
        return text.strip()

    raise ValueError("Unsupported file type. Use TXT, PDF, DOCX, JPG/PNG/BMP/TIFF.")


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        try:
            # Input 1: file or text
            if 'file1' in request.files and request.files['file1'].filename:
                text1 = parse_file(request.files['file1'])
            else:
                text1 = request.form['text1'].strip()

            # Input 2: file or text
            if 'file2' in request.files and request.files['file2'].filename:
                text2 = parse_file(request.files['file2'])
            else:
                text2 = request.form['text2'].strip()

            if not text1 or not text2:
                flash('Please provide content for both documents.')
                return render_template('index.html')

            result = detect_plagiarism(text1, text2)
            return render_template(
                'results.html',
                result=result,
                query_text=text1,
                doc_text=text2
            )
        except Exception as e:
            flash(f'Error processing files: {str(e)}')
            return render_template('index.html')

    return render_template('index.html')


if __name__ == '__main__':
    app.run(debug=True)

